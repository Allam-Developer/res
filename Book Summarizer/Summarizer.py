# pip install openai generativeai python-docx pyttsx3 moviepy PyPDF2
# add your api keys in line 260 & 261
import openai
from openai import OpenAI
import PyPDF2
import time
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
import google.generativeai as genai
import google.api_core.exceptions
import os
from moviepy.editor import concatenate_audioclips, AudioFileClip
import shutil
import pyttsx3

# setup for input
if True:
    desktop_path = os.path.expanduser("~\\Desktop\\")
    sum_list = ["1", "sum", "summ", "summarize"]
    aud_list = ["2", "aud", "audio", "mp3"]
    trns_list = ["3", "trns", "trans", "translate"]
    qna_list = ["4", "qna", "q", "mcq", "test"]
    gpt_list = ["0", "gpt", "chat gpt"]
    gemini_list = ["1", "gemini", "google"]
    agree_list = ["1", "y", "yes", "sure", "ok"]
    ar_list = ['0', 'ar' , 'arabic']
    en_list = ['1', 'en', 'english']
    male_list = ['m' , 'male' , '0' , 'man']
    female_list = ['f', 'female', '1', 'woman']
    show_list = ['show' , 'view' , 'search' , 'list' , 'all']

    # default values
    ai = '1'
    lvl = '1'
    mark_pages = '0'
    with_audio = '0'
    voice = '0'
    rate = '150'
    with_trans = '0'
    lang_target = 'ar'
    trans_ai = '1'
    with_qna = '0'
    qs_in_page = '5'
    answers = '1'
    q_ai = '1'

def manage_input(var , text):
    x = input(text)
    if not x == "":
        return x
    else:
        return var

def sleep_time(ai_name):
    if gpt_list.__contains__(ai_name):
        return 10
    elif gemini_list.__contains__(ai_name):
        return 1
def word_to_mp3(book_name, voice, rate):

    tmp_path = desktop_path + "tmp\\"
    try:
        os.makedirs(tmp_path)
        print("tmp folder created")
    except FileExistsError:
        print("already tmp folder exist")

    save_path = desktop_path + book_name.removesuffix(".docx") + ".mp3"
    book_path = desktop_path + book_name
    document = Document(book_path)


    # Initialize the engine
    engine = pyttsx3.init()
    voices = engine.getProperty('voices')
    engine.setProperty('rate', rate)
    engine.setProperty('volume', 1)
    engine.setProperty('voice', voices[int(voice)].id)

    counter = 1
    clips = []

    for paragraph in document.paragraphs:
        text = paragraph.text
        audio_path = tmp_path + "sound-" + str(counter) + ".mp3"
        engine.save_to_file(text, audio_path)
        engine.runAndWait()

        percent = str("{:.2f}".format((counter / len(document.paragraphs)) * 100)) + "%"
        print("Creating audio : " + percent)
        clips.append(AudioFileClip(audio_path))
        counter += 1

    print("combining your mp3 file")
    # combining all mp3 files created in one file
    final_clip = concatenate_audioclips(clips)
    # saving the mp3 file
    final_clip.write_audiofile(save_path)
    # delete the tmp files
    shutil.rmtree(tmp_path)

    print(book_name.removesuffix(".docx") + ".mp3")

# user input
if True:
    print("---Summarizer v0.1(Alpha)---"
          "\n1 or (sum) summarize file"
          "\n2 or (aud) create audio from file"
          "\n3 or (trns) translate file"
          "\n4 or (qna) make questions from file")

    operation = input("How Can i help You? ")

    if operation == '':
        operation = '1'

    # sum operation
    if sum_list.__contains__(operation):
        file_name = input("File Name You Want to Summarize ? ")

        # if no file provided
        if file_name == '':
            pdf_list = []
            for i in os.listdir(desktop_path):
                if i.endswith('.pdf') or i.endswith('.docx'):
                    pdf_list.append(i)
            file_name = pdf_list[0]
        if show_list.__contains__(file_name):
            pdf_list = []
            num = 1
            for i in os.listdir(desktop_path):
                if i.endswith('.pdf') or i.endswith('.docx'):
                    pdf_list.append(i)
                    print(str(num) + ' - ' + i)
                    num += 1
            file_name = pdf_list[int(input()) - 1]
        print(file_name)

        file_path = desktop_path + file_name

        if file_name.endswith('.pdf'):
            save_name = file_name.removesuffix(".pdf") + " Summary.docx"
        elif file_name.endswith('.docx'):
            save_name = file_name.removesuffix(".docx") + " Summary.docx"

        save_path = desktop_path + save_name

        ai = manage_input(ai , "Summarize With? GPT / Gemini. ")
        lvl = manage_input(lvl , "Summarizing lvl (1 --> 3) : ")
        mark_pages = manage_input(mark_pages , "mark summary with original pages (y/n)? ")

        # sum + aud operation
        with_audio = input("Create Audio Version with the Summary (y/n)? ").lower()
        if agree_list.__contains__(with_audio):
            voice = manage_input(voice , "Reader Voice Male or Female (m/f)? ")
            rate = manage_input(rate ,  "Reading Speed ? (default 150) ")

        # sum + trans operation
        with_trans = input("Create Translation with the Summary (y/n)? ").lower()
        if agree_list.__contains__(with_trans):
            lang_target = manage_input(lang_target , "Translate to (en/ar) ? ")
            trans_ai = manage_input(trans_ai , "Translate with? GPT / Gemini. ")

        # sum + qna operation
        with_qna = input("Create Mcq File with the Summary (y/n)? ").lower()
        if agree_list.__contains__(with_qna):
            qs_in_page = manage_input(qs_in_page , "Questions on each Paragraph (default 5) ? ")
            answers = manage_input(answers , "Mark Answers (y/n)? ")
            q_ai = manage_input(q_ai , "Mcq with? GPT / Gemini. ")

    # aud operation
    elif aud_list.__contains__(operation):
        file_name = input("File Name You Want to Make Audio Version for ? ")

        # if no file provided
        if file_name == '':
            pdf_list = []
            for i in os.listdir(desktop_path):
                if i.endswith('.docx'):
                    pdf_list.append(i)
            file_name = pdf_list[0]
        if show_list.__contains__(file_name):
            pdf_list = []
            num = 1
            for i in os.listdir(desktop_path):
                if i.endswith('.docx'):
                    pdf_list.append(i)
                    print(str(num) + ' - ' + i)
                    num += 1
            file_name = pdf_list[int(input()) - 1]

        voice = int(manage_input(voice , "Reader Voice Male or Female (m/f)? "))
        rate = int(manage_input(rate , "Reading Speed ? (default 150) "))

    # trans operation
    elif trns_list.__contains__(operation):
        file_name = input("File Name You Want to translate ? ")

        # if no file provided
        if file_name == '':
            pdf_list = []
            for i in os.listdir(desktop_path):
                if i.endswith('.docx'):
                    pdf_list.append(i)
            file_name = pdf_list[0]
        if show_list.__contains__(file_name):
            pdf_list = []
            num = 1
            for i in os.listdir(desktop_path):
                if i.endswith('.docx'):
                    pdf_list.append(i)
                    print(str(num) + ' - ' + i)
                    num += 1
            file_name = pdf_list[int(input()) - 1]

        lang_target = manage_input(lang_target, "Translate to (en/ar) ? ")
        trans_ai = manage_input(trans_ai, "Translate with? GPT / Gemini. ")

    # qna operation
    elif qna_list.__contains__(operation):
        file_name = input("File Name You Want Questions on ? ")

        # if no file provided
        if file_name == '':
            pdf_list = []
            for i in os.listdir(desktop_path):
                if i.endswith('.docx'):
                    pdf_list.append(i)
            file_name = pdf_list[0]
        if show_list.__contains__(file_name):
            pdf_list = []
            num = 1
            for i in os.listdir(desktop_path):
                if i.endswith('.docx'):
                    pdf_list.append(i)
                    print(str(num) + ' - ' + i)
                    num += 1
            file_name = pdf_list[int(input()) - 1]

        qs_in_page = manage_input(qs_in_page, "Questions on each Paragraph (default 5) ? ")
        answers = manage_input(answers, "Mark Answers (y/n)? ")
        q_ai = manage_input(q_ai, "Mcq with? GPT / Gemini. ")
    else:
        print('invalid input')


# setup for summarizing level
if True:

    sum_lvl1 = "Summarize this : "
    sum_lvl2 = "Summarize this as much as you can : "
    sum_lvl3 = "Summarize this in one or two sentence : "
    sum_lvls = [sum_lvl1 , sum_lvl2 , sum_lvl3]

# AI setup
if True:

    # api keys
    apiKey_gemini = "your api key here"
    apiKey_gpt = "your api key here"
    
    genai.configure(api_key=apiKey_gemini)
    client = OpenAI(api_key=apiKey_gpt)

    # Setup gemini model
    generation_config = {
        "temperature": 0.9,
        "top_p": 1,
        "top_k": 1,
        "max_output_tokens": 2048,
    }

    safety_settings = [
        {
            "category": "HARM_CATEGORY_HARASSMENT",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
        {
            "category": "HARM_CATEGORY_HATE_SPEECH",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
        {
            "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
        {
            "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
    ]

    model = genai.GenerativeModel(model_name="gemini-1.0-pro", generation_config=generation_config,
                                  safety_settings=safety_settings)
    convo = model.start_chat(history=[])

# send text to gpt and return response
def gpt_this(text):
    try:
        chat_comp = client.chat.completions.create(messages=[{"role": "user", "content": text}], model="gpt-3.5-turbo")
        response = chat_comp.choices[0].message.content
        return response
    except openai.APIConnectionError:
        chat_comp = client.chat.completions.create(messages=[{"role": "user", "content": text}], model="gpt-3.5-turbo")
        response = chat_comp.choices[0].message.content
        return response

# send text to gemini and return response
def gemini_this(text):
    connected = False
    while not connected:
        try:
            convo.send_message(text)
            response = convo.last.text
            connected = True
            return response
        except google.api_core.exceptions.InternalServerError:
            print('server error trying again')

        time.sleep(1)
# send text to AI chosen by a variable and return response
def ai_this(engine , text):
    result = ''
    if gpt_list.__contains__(engine):
        result = gpt_this(text)
    elif gemini_list.__contains__(engine):
        result = gemini_this(text)
    return result

def translate(book_name, target):
    save_path = desktop_path + book_name.removesuffix(".docx") + "-" + target + ".docx"
    book_path = desktop_path + book_name
    document_read = Document(book_path)
    document_write = Document()
    section = document_write.sections[0]

    # Set margins
    margin = 0.8
    section.top_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)

    response = ''

    # translation phrases
    ar_trans_phrase = 'translate this to arabic (translation only) : '
    en_trans_phrase = 'translate this to english (translation only) : '

    if ar_list.__contains__(target):
        response = ai_this(trans_ai, ar_trans_phrase + book_name.removesuffix(".docx"))
    elif en_list.__contains__(target):
        response = ai_this(trans_ai, en_trans_phrase + book_name.removesuffix(".docx"))

    document_write.add_heading(response, 0)
    counter = 1

    for paragraph in document_read.paragraphs:
        text = paragraph.text
        percent = str("{:.2f}".format((counter / len(document_read.paragraphs)) * 100)) + "%"
        print("Translating : " + percent)
        counter += 1
        prompt = ""
        if ar_list.__contains__(target):
            prompt = ar_trans_phrase + text
        elif en_list.__contains__(target):
            prompt = en_trans_phrase + text

        response = ai_this(trans_ai , prompt)

        para = document_write.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = para.add_run()
        if ar_list.__contains__(target):
            run.text = "\u202E" + response
        else:
            run.text = response
        if counter != len(document_read.paragraphs):
            time.sleep(sleep_time(ai))
        else:
            print("Text Done. Saving the File.")

    # save the translation file
    document_write.save(save_path)
    print(book_name.removesuffix(".docx") + "-" + target + ".docx")

def summarize():
    document = Document()
    section = document.sections[0]

    # Set margins
    margin = 0.8
    section.top_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)

    document.add_heading(save_name.removesuffix(".docx"), 0)

    # Open the pdf file
    if file_name.endswith('.pdf'):
        with open(file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Extract text from each page
            counter = 1
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()

                if text != " " and text != r"\n" and len(text) >= 40:

                    prompt = sum_lvls[int(lvl)]
                    prompt += text

                    response = ai_this(ai , prompt)

                    percent = str("{:.2f}".format((counter / len(pdf_reader.pages)) * 100)) + "%"
                    print("Summarizing " + percent + " - Pages Done : " + str(counter) + "/" + str(len(pdf_reader.pages)))
                    page_mark = " (p." + str(counter) + ")"
                    document.add_paragraph(response + page_mark)
                    counter += 1
                    if counter != len(pdf_reader.pages):
                        time.sleep(sleep_time(ai))

            print("Text Done. Saving the File.")

    # Open docx file
    elif file_name.endswith('.docx'):

        document_read = Document(file_path)
        counter = 1

        for paragraph in document_read.paragraphs:
            text = paragraph.text
            if text != " " and text != r"\n" and len(text) >= 40:

                prompt = sum_lvls[int(lvl)]
                prompt += text

                response = ai_this(ai, prompt)

                percent = str("{:.2f}".format((counter / len(document_read.paragraphs)) * 100)) + "%"
                print("Summarizing " + percent + " - Pages Done : " + str(counter) + "/" + str(len(document_read.paragraphs)))
                page_mark = " (p." + str(counter) + ")"
                document.add_paragraph(response + page_mark)
                counter += 1
                if counter != len(document_read.paragraphs):
                    time.sleep(sleep_time(ai))
                else:
                    print("Text Done. Saving the File.")

    # save the summary file
    document.save(save_path)
    print("File Saved on the Desktop With Name " + file_name.removesuffix(".pdf") + " Summary.docx")

def QNA(book_name):
    save_path = desktop_path + book_name.removesuffix(".docx") + " MCQ" + ".docx"
    book_path = desktop_path + book_name
    document_read = Document(book_path)
    document_write = Document()
    section = document_write.sections[0]

    # Set margins
    margin = 0.8
    section.top_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)

    document_write.add_heading(book_name.removesuffix(".docx") + " Mcq", 0)
    counter = 1

    for paragraph in document_read.paragraphs:
        text = paragraph.text
        percent = str("{:.2f}".format((counter / len(document_read.paragraphs)) * 100)) + "%"
        print("Making Mcq : " + percent)
        counter += 1
        if agree_list.__contains__(answers):
            prompt = 'make ' + qs_in_page + ' mcq questions on this and mark the correct answer with (*) : '
        else:
            prompt = 'make ' + qs_in_page + ' mcq questions on this dont mark the answers : '

        prompt += text
        response = ''

        if len(text) > 50:
            response = ai_this(q_ai , prompt)
            time.sleep(sleep_time(q_ai))
        document_write.add_paragraph(response)

    # save the Mcq file
    document_write.save(save_path)
    print("File Saved on the Desktop With Name " + book_name.removesuffix(".docx") + " Mcq.docx")


# do th operation
if True:
    if sum_list.__contains__(operation):
        if ai != "" and file_name != "":
            summarize()
            if agree_list.__contains__(with_audio):
                word_to_mp3(save_name, voice, rate)
            if agree_list.__contains__(with_trans):
                translate(save_name , target=lang_target)
            if agree_list.__contains__(with_qna):
                QNA(save_name)

    elif aud_list.__contains__(operation):
        word_to_mp3(file_name, voice, rate)

    elif trns_list.__contains__(operation):
        translate(file_name ,lang_target)

    elif qna_list.__contains__(operation):
        QNA(file_name)
